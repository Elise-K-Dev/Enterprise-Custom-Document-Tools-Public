from __future__ import annotations

import html
import os
import re
import secrets
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Literal
from urllib.parse import quote

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fastapi import FastAPI, HTTPException, Query, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from markdown_it import MarkdownIt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from playwright.async_api import async_playwright
from pydantic import BaseModel, Field


app = FastAPI(title="markdown-pdf-service", version="0.1.0", openapi_url=None)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


import asyncio
import logging

from app.identity import (
    ResolvedUser,
    require_admin_role,
    require_internal_request,
    resolve_user,
)
from app.grants import GrantStore, RetentionPolicy, run_cleanup_worker


OUTPUT_DIR = Path(os.getenv("MARKDOWN_PDF_OUTPUT_DIR", "/app/output"))
PUBLIC_BASE_URL = (
    os.getenv("MARKDOWN_PDF_PUBLIC_BASE_URL")
    or os.getenv("PORT_PROJECT_PUBLIC_BASE_URL")
    or "http://192.168.100.202"
).rstrip("/")
CHROMIUM_EXECUTABLE_PATH = os.getenv("CHROMIUM_EXECUTABLE_PATH", "/usr/bin/chromium")
RETENTION_CONFIG_PATH = Path(
    os.getenv("MARKDOWN_PDF_RETENTION_CONFIG", "/app/config/retention.yml")
)
GRANT_DB_PATH = Path(os.getenv("MARKDOWN_PDF_GRANT_DB", str(OUTPUT_DIR / ".grants.db")))

logging.basicConfig(
    level=os.getenv("MARKDOWN_PDF_LOG_LEVEL", "INFO"),
    format="%(asctime)s %(levelname)s %(name)s: %(message)s",
)
log = logging.getLogger("markdown_pdf")

_GRANT_STORE: GrantStore | None = None
_CLEANUP_STOP: asyncio.Event | None = None
_CLEANUP_TASK: asyncio.Task | None = None


def grant_store() -> GrantStore:
    global _GRANT_STORE
    if _GRANT_STORE is None:
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        _GRANT_STORE = GrantStore(
            db_path=GRANT_DB_PATH,
            output_dir=OUTPUT_DIR,
            config_path=RETENTION_CONFIG_PATH,
        )
    return _GRANT_STORE

REPORT_CSS = """
@page {
  size: A4;
  margin: 18mm 16mm 22mm 16mm;
}
* {
  box-sizing: border-box;
}
body {
  margin: 0;
  color: #1f2933;
  font-family: "Noto Sans CJK KR", "Noto Sans KR", "Noto Sans", Arial, sans-serif;
  font-size: 11pt;
  line-height: 1.58;
}
main {
  width: 100%;
}
h1 {
  margin: 0 0 12px;
  padding-bottom: 8px;
  border-bottom: 2px solid #1f2933;
  color: #111827;
  font-size: 22pt;
  font-weight: 700;
  line-height: 1.25;
}
h2 {
  margin: 22px 0 8px;
  color: #111827;
  font-size: 15pt;
  line-height: 1.35;
}
h3 {
  margin: 16px 0 6px;
  color: #25313f;
  font-size: 12.5pt;
}
p {
  margin: 6px 0;
}
ul, ol {
  margin: 6px 0 10px 22px;
  padding: 0;
}
li {
  margin: 3px 0;
}
table {
  width: 100%;
  margin: 10px 0 14px;
  border-collapse: collapse;
  table-layout: fixed;
  font-size: 9.5pt;
}
th, td {
  padding: 6px 7px;
  border: 1px solid #cbd5e1;
  vertical-align: top;
  word-break: keep-all;
  overflow-wrap: anywhere;
}
th {
  background: #eef2f7;
  color: #111827;
  font-weight: 700;
}
blockquote {
  margin: 10px 0;
  padding: 8px 12px;
  border-left: 4px solid #94a3b8;
  background: #f8fafc;
}
code {
  font-family: "Noto Sans Mono CJK KR", "Noto Sans Mono", monospace;
  font-size: 9.5pt;
}
pre {
  padding: 10px;
  border: 1px solid #d7dee8;
  background: #f8fafc;
  white-space: pre-wrap;
  overflow-wrap: anywhere;
}
.meta-box {
  display: grid;
  grid-template-columns: 1fr 1fr;
  gap: 4px 14px;
  margin: 0 0 16px;
  padding: 8px 10px;
  border: 1px solid #d7dee8;
  background: #f8fafc;
  color: #3e4c59;
  font-size: 9pt;
  line-height: 1.35;
}
.meta-box div {
  min-width: 0;
  overflow-wrap: anywhere;
}
.meta-label {
  color: #52606d;
  font-weight: 700;
}
.document-footer {
  position: fixed;
  left: 0;
  bottom: -13mm;
  width: 100%;
  color: #6b7280;
  font-size: 8.5pt;
  text-align: left;
}
"""


class RenderMarkdownPdfRequest(BaseModel):
    title: str = Field(..., description="보고서 제목", min_length=1, max_length=200)
    markdown: str = Field(..., description="PDF로 변환할 Markdown 본문", min_length=1)
    file_name: str | None = Field(default=None, description="생성할 PDF 파일명")
    page_size: Literal["A4", "Letter"] = Field(default="A4", description="PDF 용지 크기")
    orientation: Literal["portrait", "landscape"] = Field(default="portrait", description="PDF 방향")
    generated_for: str | None = Field(default=None, description="문서 생성 대상자 이름")
    account_name: str | None = Field(default=None, description="문서를 요청한 계정 이름")
    account_email: str | None = Field(default=None, description="문서를 요청한 계정 이메일")


class ChatMessage(BaseModel):
    role: str = Field(..., description="메시지 역할 또는 작성자 구분")
    content: str = Field(..., description="메시지 본문")
    name: str | None = Field(default=None, description="작성자 이름")
    created_at: str | None = Field(default=None, description="메시지 작성 시각")


class RenderChatDocumentRequest(BaseModel):
    title: str = Field(..., description="문서 제목", min_length=1, max_length=200)
    messages: list[ChatMessage] = Field(default_factory=list, description="내보낼 채팅 메시지 목록")
    transcript: str | None = Field(default=None, description="messages 대신 사용할 전체 본문 또는 채팅 전문")
    file_name: str | None = Field(default=None, description="생성할 문서 파일명")
    generated_for: str | None = Field(default=None, description="문서 생성 대상자 이름")
    account_name: str | None = Field(default=None, description="문서를 요청한 계정 이름")
    account_email: str | None = Field(default=None, description="문서를 요청한 계정 이메일")


class RenderMarkdownPdfResponse(BaseModel):
    output_path: str
    download_path: str
    download_url: str
    file_name: str
    title: str
    assistant_summary: str


@app.get("/health")
def health() -> dict[str, str]:
    return {"status": "ok"}


def apply_registered_user_defaults(req: Any, user: ResolvedUser) -> None:
    if not getattr(req, "generated_for", None):
        req.generated_for = user.name
    if not getattr(req, "account_name", None):
        req.account_name = user.name
    if not getattr(req, "account_email", None):
        req.account_email = user.email


def issue_grant_for(path: str, user: ResolvedUser) -> str:
    record = grant_store().issue(
        path=path,
        user_id=user.user_id,
        email=user.email,
        rank=user.rank,
    )
    return record["token"]


@app.on_event("startup")
async def _start_cleanup_worker() -> None:
    global _CLEANUP_STOP, _CLEANUP_TASK
    store = grant_store()
    _CLEANUP_STOP = asyncio.Event()
    _CLEANUP_TASK = asyncio.create_task(run_cleanup_worker(store, _CLEANUP_STOP))
    log.info(
        "grant store initialized: db=%s output=%s config=%s",
        GRANT_DB_PATH, OUTPUT_DIR, RETENTION_CONFIG_PATH,
    )


@app.on_event("shutdown")
async def _stop_cleanup_worker() -> None:
    if _CLEANUP_STOP is not None:
        _CLEANUP_STOP.set()
    if _CLEANUP_TASK is not None:
        try:
            await asyncio.wait_for(_CLEANUP_TASK, timeout=5.0)
        except Exception:
            pass


@app.get("/openapi.json")
def openapi_spec() -> dict[str, Any]:
    return {
        "openapi": "3.0.0",
        "info": {
            "title": "Document Renderer",
            "version": "0.1.0",
            "description": "Markdown 보고서를 PDF로 렌더링하고, 보고서 본문 또는 채팅 기록을 Word DOCX 또는 Excel XLSX 파일로 내보낸 뒤 다운로드 링크를 반환합니다.",
        },
        "servers": [{"url": "http://markdown-pdf-service:8003"}],
        "paths": {
            "/health": {
                "get": {
                    "operationId": "markdown_pdf_health_check",
                    "summary": "Health check",
                    "responses": {"200": {"description": "Healthy"}},
                }
            },
            "/render/markdown-pdf": {
                "post": {
                    "operationId": "render_markdown_pdf",
                    "summary": "Markdown 보고서를 PDF 파일로 생성",
                    "description": "사용자가 수리 완료 보고서, 업무 보고, 회의록, 분석 결과, 요약문을 PDF 또는 형식이 지정되지 않은 문서 파일로 요청하면 Markdown 본문을 이 도구에 전달해 PDF 다운로드 링크를 생성한다. 사용자가 Word/DOCX 또는 Excel/XLSX를 명시하면 이 도구가 아니라 해당 형식의 렌더링 도구를 호출한다. 제목은 title에만 넣고 markdown 첫 줄에 같은 제목을 반복하지 않는다. 본문은 생성 정보, 개요, 세부 내용, 표/목록, 결론 또는 조치사항 순서로 정리한다. PDF를 직접 생성할 수 없다고 답하지 말고 이 도구를 호출한다. repair_report 같은 별도 템플릿을 만들거나 create_document에 보내지 않는다. 요약과 중요도 정렬은 모델이 먼저 수행하고, 이 도구는 PDF 렌더링만 수행한다.",
                    "requestBody": {
                        "required": True,
                        "content": {
                            "application/json": {
                                "schema": {
                                    "type": "object",
                                    "required": ["title", "markdown"],
                                    "properties": {
                                        "title": {
                                            "type": "string",
                                            "description": "보고서 제목",
                                            "example": "주간 종합 업무 요약 보고서",
                                        },
                                        "markdown": {
                                            "type": "string",
                                            "description": "중요도 순으로 정리된 Markdown 보고서 본문",
                                        },
                                        "file_name": {
                                            "type": ["string", "null"],
                                            "description": "선택 PDF 파일명",
                                            "example": "weekly_summary_report.pdf",
                                        },
                                        "page_size": {
                                            "type": "string",
                                            "default": "A4",
                                            "enum": ["A4", "Letter"],
                                        },
                                        "orientation": {
                                            "type": "string",
                                            "default": "portrait",
                                            "enum": ["portrait", "landscape"],
                                        },
                                        "generated_for": {
                                            "type": ["string", "null"],
                                            "description": "문서 생성 대상자 이름. 알고 있는 현재 사용자/요청자 이름을 넣는다.",
                                        },
                                        "account_name": {
                                            "type": ["string", "null"],
                                            "description": "문서를 요청한 계정 이름",
                                        },
                                        "account_email": {
                                            "type": ["string", "null"],
                                            "description": "문서를 요청한 계정 이메일",
                                        },
                                    },
                                }
                            }
                        },
                    },
                    "responses": {"200": {"description": "PDF rendered"}},
                }
            },
            "/render/chat-docx": {
                "post": {
                    "operationId": "render_chat_docx",
                    "summary": "본문 또는 채팅 기록을 Word DOCX 파일로 내보내기",
                    "description": "사용자가 보고서, 요약문, 업무보고, 재고현황 보고서, 현재 대화 내용, 채팅 기록, 이전 답변을 워드 파일, Word 파일, DOCX 문서로 요청하면 작성한 본문을 transcript에 넣거나 messages를 전달해 DOCX 다운로드 링크를 생성한다. 제목은 title에만 넣고 transcript 첫 줄에 같은 제목을 반복하지 않는다. Word 출력에는 Markdown 문법 기호가 남지 않으며 Markdown 표는 실제 Word 표로 렌더링된다. title만 전달하지 않는다. 구매 품의서 템플릿 DOCX 생성에는 사용하지 않는다.",
                    "requestBody": {
                        "required": True,
                        "content": {
                            "application/json": {
                                "schema": {
                                    "type": "object",
                                    "required": ["title", "transcript"],
                                    "properties": {
                                        "title": {"type": "string", "description": "문서 제목"},
                                        "messages": {
                                            "type": "array",
                                            "description": "내보낼 채팅 메시지 목록",
                                            "items": {
                                                "type": "object",
                                                "required": ["role", "content"],
                                                "properties": {
                                                    "role": {"type": "string"},
                                                    "name": {"type": ["string", "null"]},
                                                    "content": {"type": "string"},
                                                    "created_at": {"type": ["string", "null"]},
                                                },
                                            },
                                        },
                                        "transcript": {
                                            "type": "string",
                                            "description": "Word 문서에 넣을 보고서 본문 또는 messages 대신 사용할 전체 채팅 전문. 빈 값으로 보내지 않는다.",
                                        },
                                        "file_name": {
                                            "type": ["string", "null"],
                                            "description": "선택 DOCX 파일명",
                                            "example": "chat_export.docx",
                                        },
                                        "generated_for": {
                                            "type": ["string", "null"],
                                            "description": "문서 생성 대상자 이름. 알고 있는 현재 사용자/요청자 이름을 넣는다.",
                                        },
                                        "account_name": {
                                            "type": ["string", "null"],
                                            "description": "문서를 요청한 계정 이름",
                                        },
                                        "account_email": {
                                            "type": ["string", "null"],
                                            "description": "문서를 요청한 계정 이메일",
                                        },
                                    },
                                }
                            }
                        },
                    },
                    "responses": {"200": {"description": "DOCX rendered"}},
                }
            },
            "/render/chat-xlsx": {
                "post": {
                    "operationId": "render_chat_xlsx",
                    "summary": "본문 또는 채팅 기록을 Excel XLSX 파일로 내보내기",
                    "description": "사용자가 보고서, 요약문, 업무보고, 재고현황 보고서, 현재 대화 내용, 채팅 기록, 이전 답변을 엑셀 파일, Excel 파일, XLSX 문서로 요청하면 표 형식 본문을 transcript에 넣거나 messages를 전달해 XLSX 다운로드 링크를 생성한다. 제목은 title에만 넣고 transcript 첫 줄에 같은 제목을 반복하지 않는다. Excel 출력에는 Markdown 문법 기호가 남지 않으며 Markdown 표는 실제 행/열로 렌더링된다. title만 전달하지 않는다.",
                    "requestBody": {
                        "required": True,
                        "content": {
                            "application/json": {
                                "schema": {
                                    "type": "object",
                                    "required": ["title", "transcript"],
                                    "properties": {
                                        "title": {"type": "string", "description": "문서 제목"},
                                        "messages": {
                                            "type": "array",
                                            "description": "내보낼 채팅 메시지 목록",
                                            "items": {
                                                "type": "object",
                                                "required": ["role", "content"],
                                                "properties": {
                                                    "role": {"type": "string"},
                                                    "name": {"type": ["string", "null"]},
                                                    "content": {"type": "string"},
                                                    "created_at": {"type": ["string", "null"]},
                                                },
                                            },
                                        },
                                        "transcript": {
                                            "type": "string",
                                            "description": "Excel 파일에 넣을 보고서 본문 또는 messages 대신 사용할 전체 채팅 전문. 빈 값으로 보내지 않는다.",
                                        },
                                        "file_name": {
                                            "type": ["string", "null"],
                                            "description": "선택 XLSX 파일명",
                                            "example": "chat_export.xlsx",
                                        },
                                        "generated_for": {
                                            "type": ["string", "null"],
                                            "description": "문서 생성 대상자 이름. 알고 있는 현재 사용자/요청자 이름을 넣는다.",
                                        },
                                        "account_name": {
                                            "type": ["string", "null"],
                                            "description": "문서를 요청한 계정 이름",
                                        },
                                        "account_email": {
                                            "type": ["string", "null"],
                                            "description": "문서를 요청한 계정 이메일",
                                        },
                                    },
                                }
                            }
                        },
                    },
                    "responses": {"200": {"description": "XLSX rendered"}},
                }
            },
            "/download": {
                "get": {
                    "operationId": "download_rendered_document_file",
                    "summary": "생성된 문서 파일 다운로드",
                    "parameters": [
                        {
                            "name": "path",
                            "in": "query",
                            "required": True,
                            "schema": {"type": "string"},
                        }
                    ],
                    "responses": {"200": {"description": "PDF file"}},
                }
            },
        },
    }


@app.post("/render/markdown-pdf", response_model=RenderMarkdownPdfResponse)
async def render_markdown_pdf(req: RenderMarkdownPdfRequest, raw_request: Request) -> RenderMarkdownPdfResponse:
    user = resolve_user(raw_request)
    apply_registered_user_defaults(req, user)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    file_name = sanitize_pdf_filename(req.file_name or req.title)
    relative_path = Path("reports") / file_name
    output_path = OUTPUT_DIR / relative_path
    output_path.parent.mkdir(parents=True, exist_ok=True)

    html_content = build_report_html(
        req.title,
        strip_duplicate_leading_title(req.markdown, req.title),
        req,
    )
    await render_pdf_with_chromium(
        html_content,
        output_path,
        page_size=req.page_size,
        landscape=req.orientation == "landscape",
    )

    token = issue_grant_for(relative_path.as_posix(), user)
    download_path = f"/download?path={quote(relative_path.as_posix())}&token={quote(token)}"
    download_url = f"{PUBLIC_BASE_URL}{download_path}"
    return RenderMarkdownPdfResponse(
        output_path=relative_path.as_posix(),
        download_path=download_path,
        download_url=download_url,
        file_name=file_name,
        title=req.title,
        assistant_summary=f"{req.title} PDF 파일을 생성했습니다. 다운로드 링크는 {download_url} 입니다.",
    )


@app.post("/render/chat-docx", response_model=RenderMarkdownPdfResponse)
async def render_chat_docx(req: RenderChatDocumentRequest, raw_request: Request) -> RenderMarkdownPdfResponse:
    user = resolve_user(raw_request)
    apply_registered_user_defaults(req, user)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    file_name = sanitize_document_filename(req.file_name or req.title, ".docx")
    relative_path = Path("reports") / file_name
    output_path = OUTPUT_DIR / relative_path
    output_path.parent.mkdir(parents=True, exist_ok=True)

    render_chat_docx_file(req, output_path)

    token = issue_grant_for(relative_path.as_posix(), user)
    download_path = f"/download?path={quote(relative_path.as_posix())}&token={quote(token)}"
    download_url = f"{PUBLIC_BASE_URL}{download_path}"
    return RenderMarkdownPdfResponse(
        output_path=relative_path.as_posix(),
        download_path=download_path,
        download_url=download_url,
        file_name=file_name,
        title=req.title,
        assistant_summary=f"{req.title} Word 파일을 생성했습니다. 다운로드 링크는 {download_url} 입니다.",
    )


@app.post("/render/chat-xlsx", response_model=RenderMarkdownPdfResponse)
async def render_chat_xlsx(req: RenderChatDocumentRequest, raw_request: Request) -> RenderMarkdownPdfResponse:
    user = resolve_user(raw_request)
    apply_registered_user_defaults(req, user)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    file_name = sanitize_document_filename(req.file_name or req.title, ".xlsx")
    relative_path = Path("reports") / file_name
    output_path = OUTPUT_DIR / relative_path
    output_path.parent.mkdir(parents=True, exist_ok=True)

    render_chat_xlsx_file(req, output_path)

    token = issue_grant_for(relative_path.as_posix(), user)
    download_path = f"/download?path={quote(relative_path.as_posix())}&token={quote(token)}"
    download_url = f"{PUBLIC_BASE_URL}{download_path}"
    return RenderMarkdownPdfResponse(
        output_path=relative_path.as_posix(),
        download_path=download_path,
        download_url=download_url,
        file_name=file_name,
        title=req.title,
        assistant_summary=f"{req.title} Excel 파일을 생성했습니다. 다운로드 링크는 {download_url} 입니다.",
    )


@app.get("/download")
def download(
    raw_request: Request,
    path: str = Query(..., description="OUTPUT_DIR 기준 상대 경로"),
    token: str | None = Query(default=None, description="등록 계정에 발급된 다운로드 토큰"),
) -> FileResponse:
    if not token:
        # Token-less internal access for legacy/tool-to-tool calls.
        resolve_user(raw_request)
    else:
        try:
            grant_store().validate_and_consume(token, path)
        except PermissionError as exc:
            raise HTTPException(status_code=403, detail=str(exc)) from exc

    requested = OUTPUT_DIR / path
    try:
        canonical_output = OUTPUT_DIR.resolve()
        canonical_file = requested.resolve()
    except FileNotFoundError as exc:
        raise HTTPException(status_code=404, detail="file not found") from exc

    if not canonical_file.is_file() or not canonical_file.is_relative_to(canonical_output):
        raise HTTPException(status_code=404, detail="file not found")

    return FileResponse(
        canonical_file,
        media_type=media_type_for_path(canonical_file),
        filename=canonical_file.name,
    )


# --- admin endpoints --------------------------------------------------------
# All require: internal token header AND hi_rank caller. Caddy already
# locks down /admin/* externally; this is the in-service guard.

def _require_admin(raw_request: Request) -> ResolvedUser:
    # require_internal_request is already enforced inside resolve_user
    user = resolve_user(raw_request)
    require_admin_role(user)
    return user


@app.get("/admin/policy")
def admin_get_policy(raw_request: Request) -> dict[str, Any]:
    _require_admin(raw_request)
    pol = grant_store().policy()
    return {"policy": pol.to_dict(), "config_path": str(RETENTION_CONFIG_PATH)}


@app.patch("/admin/policy")
async def admin_patch_policy(raw_request: Request) -> dict[str, Any]:
    _require_admin(raw_request)
    body = await raw_request.json()
    if not isinstance(body, dict):
        raise HTTPException(status_code=400, detail="body must be a JSON object")
    store = grant_store()
    applied: dict[str, Any] = {}
    for key, value in body.items():
        try:
            store.set_policy_override(key, value)
            applied[key] = value
        except ValueError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc
    return {"applied": applied, "policy": store.policy().to_dict()}


@app.get("/admin/grants")
def admin_list_grants(
    raw_request: Request,
    status: str | None = Query(default=None),
    user_id: str | None = Query(default=None),
    limit: int = Query(default=200, le=1000),
) -> dict[str, Any]:
    _require_admin(raw_request)
    return {"grants": grant_store().list_grants(status=status, user_id=user_id, limit=limit)}


@app.get("/admin/grants/{token}")
def admin_get_grant(token: str, raw_request: Request) -> dict[str, Any]:
    _require_admin(raw_request)
    row = grant_store().get(token)
    if not row:
        raise HTTPException(status_code=404, detail="grant not found")
    return row


@app.patch("/admin/grants/{token}")
async def admin_patch_grant(token: str, raw_request: Request) -> dict[str, Any]:
    _require_admin(raw_request)
    body = await raw_request.json()
    if not isinstance(body, dict):
        raise HTTPException(status_code=400, detail="body must be a JSON object")
    store = grant_store()
    try:
        return store.update_grant(
            token,
            ttl_extend_seconds=body.get("ttl_extend_seconds"),
            expires_at=body.get("expires_at"),
            max_downloads=body.get("max_downloads"),
            delete_on_download=body.get("delete_on_download"),
            status=body.get("status"),
        )
    except KeyError as exc:
        raise HTTPException(status_code=404, detail="grant not found") from exc
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc


@app.delete("/admin/grants/{token}")
def admin_delete_grant(
    token: str,
    raw_request: Request,
    delete_file: bool = Query(default=True),
) -> dict[str, Any]:
    _require_admin(raw_request)
    store = grant_store()
    row = store.get(token)
    if not row:
        raise HTTPException(status_code=404, detail="grant not found")
    store.mark_for_deletion(token)
    if delete_file:
        store._safe_unlink(OUTPUT_DIR / row["path"])
    return {"revoked": token, "path": row["path"], "file_deleted": delete_file}


@app.post("/admin/sweep")
def admin_sweep(raw_request: Request) -> dict[str, int]:
    _require_admin(raw_request)
    return grant_store().sweep()


def build_report_html(title: str, markdown: str, req: RenderMarkdownPdfRequest) -> str:
    renderer = MarkdownIt("commonmark", {"html": False}).enable("table").enable("strikethrough")
    body = renderer.render(markdown)
    meta = build_document_meta(req)
    return f"""<!doctype html>
<html lang="ko">
<head>
  <meta charset="utf-8">
  <title>{html.escape(title)}</title>
  <style>{REPORT_CSS}</style>
</head>
<body>
  <main>
    <h1>{html.escape(title)}</h1>
    {meta}
    {body}
  </main>
</body>
</html>"""


def render_chat_docx_file(req: RenderChatDocumentRequest, output_path: Path) -> None:
    rows = normalize_chat_rows(req)
    if not rows:
        raise HTTPException(status_code=400, detail="messages or transcript is required")

    try:
        doc = Document()
        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Noto Sans CJK KR"
        normal_style.font.size = Pt(10.5)
        add_docx_footer(doc, build_document_footer_text(req))

        doc.add_heading(req.title, level=0)
        doc.add_paragraph(f"내보낸 시각: {datetime.now(timezone.utc).astimezone():%Y-%m-%d %H:%M}")
        if is_transcript_only_document(req, rows):
            write_transcript_docx_content(doc, strip_duplicate_leading_title(rows[0]["content"], req.title))
            doc.save(output_path)
            return

        doc.add_paragraph(f"메시지 수: {len(rows)}")

        for idx, row in enumerate(rows, start=1):
            role = row["role"]
            name = row["name"]
            created_at = row["created_at"]
            heading = clean_markdown_text(f"{idx}. {display_speaker(role, name)}")
            if created_at:
                heading = f"{heading} ({clean_markdown_text(created_at)})"
            doc.add_heading(heading, level=2)
            for paragraph in split_message_paragraphs(row["content"]):
                doc.add_paragraph(clean_markdown_text(paragraph))

        doc.save(output_path)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"failed to render chat DOCX: {exc}") from exc


def render_chat_xlsx_file(req: RenderChatDocumentRequest, output_path: Path) -> None:
    rows = normalize_chat_rows(req)
    if not rows:
        raise HTTPException(status_code=400, detail="messages or transcript is required")

    try:
        wb = Workbook()
        sheet = wb.active
        sheet.title = "채팅기록"
        sheet.oddFooter.left.text = build_document_footer_text(req)
        sheet.append(["제목", req.title])
        sheet.append(["내보낸 시각", f"{datetime.now(timezone.utc).astimezone():%Y-%m-%d %H:%M}"])
        sheet.append(["메시지 수", len(rows)])
        sheet.append([])
        transcript_content = strip_duplicate_leading_title(rows[0]["content"], req.title) if is_transcript_only_document(req, rows) else ""
        if is_transcript_only_document(req, rows) and append_transcript_tables_to_sheet(sheet, transcript_content):
            header_row = 5
        else:
            sheet.append(["순번", "역할", "작성자", "시각", "내용"])
            for idx, row in enumerate(rows, start=1):
                sheet.append(
                    [
                        idx,
                        clean_markdown_text(row["role"]),
                        clean_markdown_text(row["name"]),
                        clean_markdown_text(row["created_at"]),
                        clean_markdown_document_text(strip_duplicate_leading_title(row["content"], req.title)),
                    ]
                )
            header_row = 5

        for cell in sheet[header_row]:
            cell.font = Font(bold=True, color="111827")
            cell.fill = PatternFill("solid", fgColor="EEF2F7")
        for row in sheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
        sheet.freeze_panes = "A6"
        autosize_columns(sheet)
        if not (is_transcript_only_document(req, rows) and sheet.max_column > 5):
            sheet.column_dimensions["E"].width = 80
        wb.save(output_path)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"failed to render chat XLSX: {exc}") from exc


def normalize_chat_rows(req: RenderChatDocumentRequest) -> list[dict[str, str]]:
    rows = [
        {
            "role": message.role.strip() or "message",
            "name": (message.name or "").strip(),
            "created_at": (message.created_at or "").strip(),
            "content": message.content.strip(),
        }
        for message in req.messages
        if message.content.strip()
    ]
    if rows:
        return rows

    transcript = (req.transcript or "").strip()
    if not transcript:
        return []
    return [
        {
            "role": "transcript",
            "name": "",
            "created_at": "",
            "content": transcript,
        }
    ]


def is_transcript_only_document(req: RenderChatDocumentRequest, rows: list[dict[str, str]]) -> bool:
    has_message_content = any(message.content.strip() for message in req.messages)
    return not has_message_content and len(rows) == 1 and rows[0]["role"] == "transcript"


def build_document_footer_text(req: Any) -> str:
    generated_for = clean_footer_value(getattr(req, "generated_for", None)) or "미지정"
    account_name = clean_footer_value(getattr(req, "account_name", None))
    account_email = clean_footer_value(getattr(req, "account_email", None))
    account = "미지정"
    if account_name and account_email:
        account = f"{account_name} <{account_email}>"
    elif account_name:
        account = account_name
    elif account_email:
        account = account_email
    generated_at = datetime.now(timezone.utc).astimezone().strftime("%Y-%m-%d %H:%M")
    return f"대상자: {generated_for} | 생성시각: {generated_at} | 계정: {account}"


def build_document_meta(req: Any) -> str:
    generated_for = clean_footer_value(getattr(req, "generated_for", None)) or "미지정"
    account_name = clean_footer_value(getattr(req, "account_name", None))
    account_email = clean_footer_value(getattr(req, "account_email", None))
    account = "미지정"
    if account_name and account_email:
        account = f"{account_name} <{account_email}>"
    elif account_name:
        account = account_name
    elif account_email:
        account = account_email
    generated_at = datetime.now(timezone.utc).astimezone().strftime("%Y-%m-%d %H:%M")

    rows = [
        ("대상자", generated_for),
        ("생성 시각", generated_at),
        ("계정", account),
    ]
    items = "\n".join(
        f'<div><span class="meta-label">{html.escape(label)}</span>: {html.escape(value)}</div>'
        for label, value in rows
    )
    return f'<section class="meta-box">{items}</section>'


def clean_footer_value(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def add_docx_footer(doc: Any, footer_text: str) -> None:
    for section in doc.sections:
        paragraph = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        paragraph.text = footer_text
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.size = Pt(8)


def strip_duplicate_leading_title(content: str, title: str) -> str:
    lines = (content or "").splitlines()
    while lines and not lines[0].strip():
        lines.pop(0)
    if not lines:
        return content

    first = lines[0].strip()
    first = re.sub(r"^#{1,6}\s+", "", first)
    first = clean_markdown_text(first)
    if normalize_title_text(first) == normalize_title_text(title):
        return "\n".join(lines[1:]).lstrip()
    return content


def normalize_title_text(value: str) -> str:
    return re.sub(r"[\s_\-:：|]+", "", clean_markdown_text(value or "")).casefold()


def write_transcript_docx_content(doc: Any, content: str) -> None:
    lines = [line.rstrip() for line in content.splitlines()]
    index = 0
    while index < len(lines):
        raw_line = lines[index]
        stripped = raw_line.strip()
        if not stripped:
            index += 1
            continue

        table_rows, next_index = parse_markdown_table(lines, index)
        if table_rows:
            add_docx_table(doc, table_rows)
            index = next_index
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading_match:
            doc.add_heading(
                clean_markdown_text(heading_match.group(2)),
                level=len(heading_match.group(1)),
            )
            index += 1
            continue

        bullet_match = re.match(r"^[-*]\s+(.+)$", stripped)
        numbered_match = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if bullet_match:
            doc.add_paragraph(clean_markdown_text(bullet_match.group(1)), style="List Bullet")
        elif numbered_match:
            doc.add_paragraph(clean_markdown_text(numbered_match.group(1)), style="List Number")
        else:
            doc.add_paragraph(clean_markdown_text(stripped))
        index += 1


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]] | None, int]:
    if start + 1 >= len(lines):
        return None, start
    header_line = lines[start].strip()
    separator_line = lines[start + 1].strip()
    if not is_table_row(header_line) or not is_table_separator(separator_line):
        return None, start

    rows = [split_markdown_table_row(header_line)]
    index = start + 2
    while index < len(lines) and is_table_row(lines[index].strip()):
        row = split_markdown_table_row(lines[index].strip())
        if len(row) < 2:
            break
        rows.append(row)
        index += 1

    column_count = max(len(row) for row in rows)
    normalized_rows = [
        [clean_markdown_text(cell) for cell in row + [""] * (column_count - len(row))]
        for row in rows
    ]
    return normalized_rows, index


def is_table_row(line: str) -> bool:
    return "|" in line and len(split_markdown_table_row(line)) >= 2


def is_table_separator(line: str) -> bool:
    cells = split_markdown_table_row(line)
    if len(cells) < 2:
        return False
    return all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def split_markdown_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def add_docx_table(doc: Any, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, row_values in enumerate(rows):
        for col_index, value in enumerate(row_values):
            cell = table.cell(row_index, col_index)
            cell.text = value
            if row_index == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    doc.add_paragraph()


def append_transcript_tables_to_sheet(sheet: Any, content: str) -> bool:
    lines = [line.rstrip() for line in content.splitlines()]
    index = 0
    appended = False
    while index < len(lines):
        table_rows, next_index = parse_markdown_table(lines, index)
        if table_rows:
            for row in table_rows:
                sheet.append(row)
            sheet.append([])
            appended = True
            index = next_index
        else:
            index += 1
    if appended and sheet.max_row > 0 and not any(cell.value for cell in sheet[sheet.max_row]):
        sheet.delete_rows(sheet.max_row)
    return appended


def clean_markdown_text(text: str) -> str:
    cleaned = text or ""
    cleaned = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", cleaned)
    cleaned = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", cleaned)
    cleaned = re.sub(r"`([^`]+)`", r"\1", cleaned)
    cleaned = re.sub(r"~~(.+?)~~", r"\1", cleaned)
    cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"__(.+?)__", r"\1", cleaned)
    cleaned = re.sub(r"(?<!\w)\*(?!\s)(.+?)(?<!\s)\*(?!\w)", r"\1", cleaned)
    cleaned = re.sub(r"(?<!\w)_(?!\s)(.+?)(?<!\s)_(?!\w)", r"\1", cleaned)
    cleaned = re.sub(r"^#{1,6}\s+", "", cleaned.strip())
    return cleaned.replace("\\|", "|").strip()


def clean_markdown_document_text(text: str) -> str:
    cleaned_lines: list[str] = []
    for line in (text or "").splitlines():
        stripped = line.strip()
        if not stripped:
            cleaned_lines.append("")
            continue
        if is_table_separator(stripped):
            continue
        stripped = re.sub(r"^#{1,6}\s+", "", stripped)
        stripped = re.sub(r"^[-*]\s+", "", stripped)
        stripped = re.sub(r"^\d+[.)]\s+", "", stripped)
        cleaned_lines.append(clean_markdown_text(stripped))
    return "\n".join(cleaned_lines).strip()


def display_speaker(role: str, name: str) -> str:
    if name:
        return f"{name} [{role}]"
    return role


def split_message_paragraphs(content: str) -> list[str]:
    paragraphs = [part.strip() for part in re.split(r"\n{2,}", content) if part.strip()]
    return paragraphs or [content]


def autosize_columns(sheet: Any) -> None:
    for column in sheet.columns:
        max_length = 8
        column_letter = get_column_letter(column[0].column)
        for cell in column:
            value = "" if cell.value is None else str(cell.value)
            max_length = max(max_length, min(len(value), 60))
        sheet.column_dimensions[column_letter].width = max_length + 2


async def render_pdf_with_chromium(
    html_content: str,
    output_path: Path,
    page_size: str,
    landscape: bool,
) -> None:
    try:
        async with async_playwright() as playwright:
            chromium_args = ["--disable-dev-shm-usage"]
            if os.getenv("CHROMIUM_DISABLE_SANDBOX", "false").lower() == "true":
                chromium_args.append("--no-sandbox")
            browser = await playwright.chromium.launch(
                headless=True,
                executable_path=CHROMIUM_EXECUTABLE_PATH,
                args=chromium_args,
            )
            page = await browser.new_page()
            await page.set_content(html_content, wait_until="networkidle")
            await page.pdf(
                path=str(output_path),
                format=page_size,
                landscape=landscape,
                print_background=True,
                margin={"top": "0", "right": "0", "bottom": "0", "left": "0"},
            )
            await browser.close()
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"failed to render PDF: {exc}") from exc


def sanitize_pdf_filename(raw: str) -> str:
    return sanitize_document_filename(raw, ".pdf")


def sanitize_document_filename(raw: str, suffix: str) -> str:
    base = raw.rsplit("/", 1)[-1].rsplit("\\", 1)[-1].strip()
    if base.lower().endswith(suffix):
        base = base[: -len(suffix)]
    base = re.sub(r"[^0-9A-Za-z가-힣._ -]+", "_", base)
    base = re.sub(r"\s+", "_", base).strip("._- ")
    if not base:
        base = "report"
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{base}_{stamp}{suffix}"


def media_type_for_path(path: Path) -> str:
    extension = path.suffix.lower()
    if extension == ".pdf":
        return "application/pdf"
    if extension == ".docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if extension == ".xlsx":
        return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    return "application/octet-stream"
